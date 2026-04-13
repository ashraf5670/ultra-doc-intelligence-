"""
Ultra Doc-Intelligence — FastAPI Backend
RAG pipeline for logistics document Q&A (Gemini)
"""

import os
import uuid
import json
import io
import re
from pathlib import Path

import numpy as np
import faiss
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel

from google import genai
import pdfplumber
import docx as python_docx
from sentence_transformers import SentenceTransformer

# ── Load .env file ────────────────────────────────────────────────────────────
env_path = Path(__file__).parent.parent / ".env"
if env_path.exists():
    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())

# ── App ──────────────────────────────────────────────────────────────────────
app = FastAPI(title="Ultra Doc-Intelligence", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Models & Clients ─────────────────────────────────────────────────────────
embedder = SentenceTransformer("all-MiniLM-L6-v2")

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not set. Please add it to your .env file.")

gemini = genai.Client(api_key=GEMINI_API_KEY)
MODEL  = "gemini-2.0-flash-lite"

# ── In-Memory Store ───────────────────────────────────────────────────────────
doc_store: dict = {}

# ── Schemas ───────────────────────────────────────────────────────────────────
class AskRequest(BaseModel):
    doc_id: str
    question: str
    top_k: int = 5
    confidence_threshold: float = 0.35

class ExtractRequest(BaseModel):
    doc_id: str

# ── Helpers ───────────────────────────────────────────────────────────────────

def parse_document(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)

    elif ext == ".docx":
        doc = python_docx.Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    parts.append(" | ".join(row_texts))
        return "\n".join(parts)

    elif ext in (".txt", ".text"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 300, overlap: int = 60) -> list[str]:
    words = text.split()
    chunks, start = [], 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        start += chunk_size - overlap
    return [c for c in chunks if len(c.strip()) > 20]


def embed(texts: list[str]) -> np.ndarray:
    vecs = embedder.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    norms = np.linalg.norm(vecs, axis=1, keepdims=True)
    return (vecs / np.maximum(norms, 1e-9)).astype("float32")


def build_index(embeddings: np.ndarray) -> faiss.IndexFlatIP:
    dim   = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)
    return index


def retrieve(doc_id: str, question: str, top_k: int):
    store = doc_store[doc_id]
    q_vec = embed([question])
    scores, idxs = store["index"].search(q_vec, min(top_k, len(store["chunks"])))
    chunks = [store["chunks"][i] for i in idxs[0]]
    return chunks, scores[0].tolist()


def confidence_score(scores: list[float]) -> float:
    if not scores:
        return 0.0
    top1 = float(scores[0])
    mean = float(np.mean(scores))
    raw  = 0.65 * top1 + 0.35 * mean
    return round(min(max(raw, 0.0), 1.0), 4)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    allowed = {".pdf", ".docx", ".txt", ".text"}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed:
        raise HTTPException(400, f"Unsupported type '{ext}'. Allowed: {allowed}")

    raw = await file.read()
    try:
        text = parse_document(raw, file.filename)
    except Exception as e:
        raise HTTPException(422, f"Parse error: {e}")

    if len(text.strip()) < 50:
        raise HTTPException(422, "Document appears empty or unreadable.")

    chunks     = chunk_text(text)
    embeddings = embed(chunks)
    index      = build_index(embeddings)
    doc_id     = str(uuid.uuid4())

    doc_store[doc_id] = {
        "filename":   file.filename,
        "text":       text,
        "chunks":     chunks,
        "embeddings": embeddings,
        "index":      index,
    }

    return {
        "doc_id":      doc_id,
        "filename":    file.filename,
        "num_chunks":  len(chunks),
        "char_count":  len(text),
    }


@app.post("/ask")
async def ask(req: AskRequest):
    if req.doc_id not in doc_store:
        raise HTTPException(404, "Document not found. Upload first.")

    chunks, scores = retrieve(req.doc_id, req.question, req.top_k)
    conf = confidence_score(scores)

    if conf < req.confidence_threshold:
        return {
            "answer":      "Not found in document.",
            "confidence":  conf,
            "sources":     [],
            "guardrail":   "low_confidence",
            "explanation": f"Similarity {scores[0]:.3f} below threshold {req.confidence_threshold}.",
        }

    context = "\n\n---\n\n".join(
        f"[Chunk {i+1}]\n{c}" for i, c in enumerate(chunks)
    )

    prompt = (
        "You are a logistics document assistant. "
        "Answer ONLY from the provided document context. "
        "If the answer is not present in the context, say exactly: "
        "Not found in document.\n"
        "Be concise and precise.\n\n"
        f"Context:\n\n{context}\n\n"
        f"Question: {req.question}\n\nAnswer:"
    )

    try:
        response = gemini.models.generate_content(model=MODEL, contents=prompt)
        answer = response.text.strip()
    except Exception as e:
        raise HTTPException(500, f"Gemini error: {str(e)}")

    if "not found in document" in answer.lower():
        return {
            "answer":     "Not found in document.",
            "confidence": conf,
            "sources":    [],
            "guardrail":  "model_declined",
        }

    return {
        "answer":     answer,
        "confidence": conf,
        "sources":    chunks[:3],
        "guardrail":  None,
    }


@app.post("/extract")
async def extract(req: ExtractRequest):
    if req.doc_id not in doc_store:
        raise HTTPException(404, "Document not found. Upload first.")

    text    = doc_store[req.doc_id]["text"]
    excerpt = text[:6000]

    prompt = (
        "You are a logistics data extraction engine. "
        "Extract the following fields from the document. "
        "Return ONLY a valid JSON object — no markdown, no explanation. "
        "Use null for any field not present in the document. "
        "Fields: shipment_id, shipper, consignee, pickup_datetime, "
        "delivery_datetime, equipment_type, mode, rate, currency, "
        "weight, carrier_name\n\n"
        f"Document:\n\n{excerpt}"
    )

    try:
        response = gemini.models.generate_content(model=MODEL, contents=prompt)
        raw = response.text.strip()
    except Exception as e:
        raise HTTPException(500, f"Gemini error: {str(e)}")

    raw = re.sub(r"^```[a-z]*\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        data = {"raw_output": raw, "parse_error": "LLM returned non-JSON"}

    return {"doc_id": req.doc_id, "extraction": data}


@app.get("/health")
def health():
    return {"status": "ok", "documents_loaded": len(doc_store)}


# ── Serve frontend ────────────────────────────────────────────────────────────
frontend_dir = Path(__file__).parent.parent / "frontend"
if frontend_dir.exists():
    app.mount("/static", StaticFiles(directory=str(frontend_dir)), name="static")

    @app.get("/")
    def root():
        return FileResponse(str(frontend_dir / "index.html"))